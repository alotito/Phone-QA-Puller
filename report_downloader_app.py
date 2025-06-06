# report_downloader_app.py

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import configparser
import os
import sys
from typing import Dict, Any, Optional

# --- Dependency Imports ---
try:
    import pyodbc
    from docx import Document
    from docx.shared import Inches, RGBColor
except ImportError as e:
    missing_module = str(e).split("'")[1]
    messagebox.showerror(
        "Missing Dependency",
        f"A required library is missing: '{missing_module}'.\n\nPlease install it by running:\npip install {missing_module}"
    )
    sys.exit(1)

# --- Configuration ---
CONFIG_FILE_NAME = "config.ini"

# --- DOCX Generation Logic (borrowed from AutoQA.py for consistency) ---
def _apply_finding_color_to_run(run, finding_text):
    if not finding_text: return
    if finding_text.lower() == "positive":
        run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
    elif finding_text.lower() == "negative":
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif finding_text.lower() == "neutral":
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x8B)
    run.bold = True

def generate_docx_from_combined_json(json_data: Dict[str, Any], docx_filepath: str):
    """Generates a .docx file from a combined analysis data dictionary."""
    try:
        doc = Document()
        header = json_data.get("report_header", {})
        agent_display_name = header.get('agent_name', 'N/A')
        
        doc.add_heading(f"Performance Trend Analysis & Coaching Report: {agent_display_name}", level=1)
        doc.add_paragraph(f"Analysis based on {header.get('number_of_reports_successfully_analyzed', 'N/A')} calls. Period: {header.get('analysis_period_note', 'N/A')}")

        qualitative_summary = json_data.get("qualitative_summary_and_coaching_plan", {})
        
        if strengths := qualitative_summary.get("overall_strengths_observed", []):
            doc.add_heading("Key Strengths & Consistent Positive Performance", level=2)
            for item in strengths:
                doc.add_paragraph(item, style='ListBullet')

        if dev_areas := qualitative_summary.get("overall_areas_for_development", []):
            doc.add_heading("Areas for Development & Recurring Challenges", level=2)
            for item in dev_areas:
                doc.add_paragraph(item, style='ListBullet')

        if coaching_items := qualitative_summary.get("consolidated_coaching_focus", []):
            doc.add_heading("Consolidated Coaching & Development Plan", level=2)
            for area_item in coaching_items:
                doc.add_heading(area_item.get("area", "Focus Area"), level=3)
                if recommendations := area_item.get("specific_actions", []):
                    doc.add_paragraph().add_run("Actionable Recommendations:").bold = True
                    for rec_item in recommendations:
                        doc.add_paragraph(rec_item, style='ListNumber')
        
        if detailed_qpa := json_data.get("detailed_quality_point_analysis", []):
            doc.add_heading("Detailed Quality Point Analysis", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Quality Point"
            hdr_cells[1].text = "Trend Observation"
            hdr_cells[2].text = "Coaching Recommendation"
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            for item in detailed_qpa:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get("quality_point", "N/A")
                row_cells[1].text = item.get("trend_observation", "N/A")
                row_cells[2].text = item.get("coaching_recommendation_for_point", "N/A")

        doc.save(docx_filepath)
        return True
    except Exception as e:
        messagebox.showerror("DOCX Generation Error", f"Failed to generate the .docx file:\n{e}")
        return False

# --- Main Application Class ---
class ReportDownloaderApp:
    def __init__(self, master):
        self.master = master
        master.title("QA Report Downloader")
        master.geometry("600x450")

        self.conn = self.get_db_connection()
        if not self.conn:
            master.destroy()
            return

        # --- Data Storage ---
        self.agents = {}  # {agent_name: agent_id}
        self.analysis_dates = {}  # {formatted_date_str: analysis_id}

        # --- UI Elements ---
        # Agent Selection
        ttk.Label(master, text="1. Select an Agent:", font=("Segoe UI", 10, "bold")).pack(pady=(10, 2), anchor="w", padx=10)
        self.agent_listbox = tk.Listbox(master, exportselection=False, height=8)
        self.agent_listbox.pack(fill="x", expand=True, padx=10)
        self.agent_listbox.bind("<<ListboxSelect>>", self.on_agent_select)

        # Date Selection
        ttk.Label(master, text="2. Select a Report Date:", font=("Segoe UI", 10, "bold")).pack(pady=(10, 2), anchor="w", padx=10)
        self.date_listbox = tk.Listbox(master, exportselection=False, height=8)
        self.date_listbox.pack(fill="x", expand=True, padx=10)

        # Download Button
        self.download_button = ttk.Button(master, text="Download Selected Report as .DOCX", command=self.on_download_click)
        self.download_button.pack(pady=15, padx=10, fill="x")

        self.populate_agent_list()

    def get_db_connection(self) -> Optional[pyodbc.Connection]:
        """Reads config and establishes a database connection."""
        try:
            config_path = os.path.join(getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__))), CONFIG_FILE_NAME)
            if not os.path.exists(config_path):
                messagebox.showerror("Configuration Error", f"Config file '{CONFIG_FILE_NAME}' not found.")
                return None
            
            config = configparser.ConfigParser()
            config.read(config_path)
            db_config = config['Database']
            conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={db_config['Server']};DATABASE={db_config['Database']};UID={db_config['User']};PWD={db_config['Password']};"
            return pyodbc.connect(conn_str)
        except Exception as e:
            messagebox.showerror("Database Connection Error", f"Could not connect to the database:\n{e}")
            return None

    def populate_agent_list(self):
        """Fetches agents from DB and populates the listbox."""
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT AgentID, AgentName FROM Agents ORDER BY AgentName;")
            for row in cursor.fetchall():
                self.agents[row.AgentName] = row.AgentID
                self.agent_listbox.insert(tk.END, row.AgentName)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch agent list:\n{e}")

    def on_agent_select(self, event=None):
        """Handler for when an agent is selected."""
        selections = self.agent_listbox.curselection()
        if not selections:
            return

        # Clear previous entries
        self.date_listbox.delete(0, tk.END)
        self.analysis_dates.clear()

        selected_agent_name = self.agent_listbox.get(selections[0])
        agent_id = self.agents.get(selected_agent_name)

        if agent_id:
            try:
                cursor = self.conn.cursor()
                sql = "SELECT CombinedAnalysisID, AnalysisDateTime FROM CombinedAnalyses WHERE AgentID = ? ORDER BY AnalysisDateTime DESC;"
                cursor.execute(sql, agent_id)
                for row in cursor.fetchall():
                    # Format date for display
                    date_str = row.AnalysisDateTime.strftime("%Y-%m-%d %I:%M %p")
                    self.analysis_dates[date_str] = row.CombinedAnalysisID
                    self.date_listbox.insert(tk.END, date_str)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to fetch reports for {selected_agent_name}:\n{e}")

    def on_download_click(self):
        """Handler for the download button click."""
        date_selections = self.date_listbox.curselection()
        agent_selections = self.agent_listbox.curselection()
        
        if not agent_selections:
            messagebox.showwarning("Selection Required", "Please select an agent first.")
            return
        if not date_selections:
            messagebox.showwarning("Selection Required", "Please select a report date.")
            return

        selected_date_str = self.date_listbox.get(date_selections[0])
        analysis_id = self.analysis_dates.get(selected_date_str)
        agent_name = self.agent_listbox.get(agent_selections[0])

        if analysis_id:
            report_data = self.fetch_combined_analysis_data(analysis_id)
            if not report_data:
                messagebox.showerror("Data Error", "Could not retrieve the report data from the database.")
                return

            # Ask user for save location
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                title="Save Report As...",
                initialfile=f"QA_Report_{agent_name}_{selected_date_str.split(' ')[0]}.docx"
            )
            
            if save_path:
                if generate_docx_from_combined_json(report_data, save_path):
                    messagebox.showinfo("Success", f"Report successfully saved to:\n{save_path}")

    def fetch_combined_analysis_data(self, analysis_id: int) -> Optional[Dict]:
        """Queries all related tables and reconstructs the report data dictionary."""
        try:
            cursor = self.conn.cursor()
            
            # Main analysis data
            main_sql = "SELECT * FROM CombinedAnalyses c JOIN Agents a ON c.AgentID = a.AgentID WHERE c.CombinedAnalysisID = ?;"
            main_row = cursor.execute(main_sql, analysis_id).fetchone()
            if not main_row: return None

            report = {
                "report_header": {
                    "agent_name": main_row.AgentName,
                    "number_of_reports_successfully_analyzed": main_row.NumberOfReportsSuccessfullyAnalyzed,
                    "analysis_period_note": main_row.AnalysisPeriodNote
                },
                "qualitative_summary_and_coaching_plan": {
                    "overall_strengths_observed": [row.StrengthText for row in cursor.execute("SELECT StrengthText FROM CombinedAnalysisStrengths WHERE CombinedAnalysisID = ?", analysis_id).fetchall()],
                    "overall_areas_for_development": [row.DevelopmentAreaText for row in cursor.execute("SELECT DevelopmentAreaText FROM CombinedAnalysisDevelopmentAreas WHERE CombinedAnalysisID = ?", analysis_id).fetchall()],
                    "consolidated_coaching_focus": []
                },
                "detailed_quality_point_analysis": []
            }
            
            # Coaching Focus
            focus_rows = cursor.execute("SELECT CoachingFocusID, AreaText FROM CombinedAnalysisCoachingFocus WHERE CombinedAnalysisID = ?", analysis_id).fetchall()
            for focus_row in focus_rows:
                actions = [row.ActionText for row in cursor.execute("SELECT ActionText FROM CombinedAnalysisCoachingActions WHERE CoachingFocusID = ?", focus_row.CoachingFocusID).fetchall()]
                report["qualitative_summary_and_coaching_plan"]["consolidated_coaching_focus"].append({
                    "area": focus_row.AreaText,
                    "specific_actions": actions
                })

            # Detailed QP Analysis
            detail_sql = """
                SELECT qp.QualityPointText, d.TrendObservation
                FROM CombinedAnalysisQualityPointDetails d
                JOIN QualityPointsMaster qp ON d.QualityPointID = qp.QualityPointID
                WHERE d.CombinedAnalysisID = ?;
            """
            for detail_row in cursor.execute(detail_sql, analysis_id).fetchall():
                report["detailed_quality_point_analysis"].append({
                    "quality_point": detail_row.QualityPointText,
                    "trend_observation": detail_row.TrendObservation,
                    "coaching_recommendation_for_point": "" # Note: This column wasn't in the final schema, add if needed
                })

            return report
        except Exception as e:
            messagebox.showerror("Database Query Error", f"An error occurred while fetching report details:\n{e}")
            return None

    def on_closing(self):
        """Handle window closing event."""
        if self.conn:
            self.conn.close()
        self.master.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = ReportDownloaderApp(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()
